import sys
import os
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.uic import loadUi
from reportlab.pdfbase.pdfmetrics import registerFontFamily
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.lib import colors
import tempfile
registerFontFamily('DejaVuSerif', normal='DejaVuSerif',
                   bold='DejaVuSerif-Bold', italic='DejaVuSerif-Italic')
outpath = os.path.join(os.getcwd(), 'font')
pdfmetrics.registerFont(TTFont('DejaVuSans', os.path.join(outpath, 'DejaVuSans.ttf')))
pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', os.path.join(outpath, 'DejaVuSans-Bold.ttf')))
pdfmetrics.registerFont(TTFont('DejaVuSerif', os.path.join(outpath, 'DejaVuSerif.ttf')))
pdfmetrics.registerFont(TTFont('DejaVuSerif-Bold', os.path.join(outpath, 'DejaVuSerif-Bold.ttf')))
pdfmetrics.registerFont(TTFont('DejaVuSerif-Italic', os.path.join(outpath, 'DejaVuSerif-Italic.ttf')))
class ProtocolGenerator(QMainWindow):
    def __init__(self):
        super().__init__()
        
        # Настройка главного окна
        self.setWindowTitle("Генератор протоколов командного матча")
        self.setGeometry(100, 100, 1200, 800)
        
        # Центральный виджет
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # Прокручиваемая область
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        main_layout.addWidget(scroll_area)
        
        # Контейнер для протоколов
        container = QWidget()
        scroll_area.setWidget(container)
        self.protocols_layout = QVBoxLayout(container)
        self.protocols_layout.setSpacing(20)
        
        # Создаем 4 протокола
        self.protocol_widgets = []
        for i in range(4):
            protocol = self.create_protocol_widget(i+1)
            self.protocols_layout.addWidget(protocol)
            self.protocol_widgets.append(protocol)
        
        # Кнопки управления
        button_layout = QHBoxLayout()
        
        self.btn_add_data = QPushButton("Заполнить тестовыми данными")
        self.btn_add_data.clicked.connect(self.fill_test_data)
        button_layout.addWidget(self.btn_add_data)
        
        self.btn_generate_pdf = QPushButton("Создать PDF (4 на листе)")
        # self.btn_generate_pdf.clicked.connect(self.save_to_word)
        self.btn_generate_pdf.clicked.connect(self.generate_pdf)
        button_layout.addWidget(self.btn_generate_pdf)
        
        self.btn_print = QPushButton("Печать")
        # self.btn_print.clicked.connect(self.print_protocols)
        self.btn_print.clicked.connect(self.create_exact_table_widget)
        button_layout.addWidget(self.btn_print)
        
        self.btn_clear = QPushButton("Очистить все")
        self.btn_clear.clicked.connect(self.clear_all)
        button_layout.addWidget(self.btn_clear)
        
        main_layout.addLayout(button_layout)
        
        # Хранилище данных
        self.protocol_data = []
        
    def create_protocol_widget(self, protocol_num):
        """Создание виджета для одного протокола"""
        protocol_frame = QFrame()
        protocol_frame.setFrameStyle(QFrame.Box | QFrame.Raised)
        protocol_frame.setLineWidth(2)
        layout = QVBoxLayout(protocol_frame)
        layout.setSpacing(5)
        
        # Заголовок протокола
        header_label = QLabel(f"ПРОТОКОЛ #{protocol_num}")
        header_label.setAlignment(Qt.AlignCenter)
        header_label.setStyleSheet("""
            QLabel {
                font-weight: bold;
                font-size: 14pt;
                color: #000080;
                padding: 5px;
                background-color: #f0f0f0;
                border: 1px solid #808080;
            }
        """)
        layout.addWidget(header_label)
        
        # Основной контейнер для данных протокола
        data_container = QWidget()
        data_layout = QVBoxLayout(data_container)
        data_layout.setSpacing(10)
        
        # 1. Шапка документа (как в Word)
        header_widget = self.create_header_widget()
        data_layout.addWidget(header_widget)
        
        # 2. Основная таблица
        table_widget = self.create_main_table()
        data_layout.addWidget(table_widget)
        
        # 3. Подписи
        signatures_widget = self.create_signatures_widget()
        data_layout.addWidget(signatures_widget)
        
        # 4. Легенда сокращений
        legend_widget = self.create_legend_widget()
        data_layout.addWidget(legend_widget)
        
        # Прокрутка для контейнера данных
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setWidget(data_container)
        scroll_area.setMinimumHeight(400)
        scroll_area.setMaximumHeight(600)
        
        layout.addWidget(scroll_area)
        
        # Сохраняем ссылки на важные виджеты
        protocol_frame.table_widget = table_widget
        protocol_frame.signatures = signatures_widget
        
        return protocol_frame
    
    def create_header_widget(self):
        """Создание шапки документа"""
        header_widget = QWidget()
        layout = QVBoxLayout(header_widget)
        layout.setSpacing(2)
        
        # Министерство и Федерация
        ministry_label = QLabel("Министерство спорта Российской Федерации / Федерация настольного тенниса России")
        ministry_label.setAlignment(Qt.AlignCenter)
        ministry_label.setStyleSheet("font-weight: bold; font-size: 10pt;")
        layout.addWidget(ministry_label)
        
        # Директорат
        directorate_label = QLabel("Директорат Командного Чемпионата ФНТР")
        directorate_label.setAlignment(Qt.AlignCenter)
        directorate_label.setStyleSheet("font-weight: bold; font-size: 10pt;")
        layout.addWidget(directorate_label)
        
        # Основной заголовок
        main_header = QLabel("ПРОТОКОЛ КОМАНДНОГО МАТЧА")
        main_header.setAlignment(Qt.AlignCenter)
        main_header.setStyleSheet("font-weight: bold; font-size: 14pt; text-decoration: underline;")
        layout.addWidget(main_header)
        
        # Чемпионат
        championship_label = QLabel("КОМАНДНОГО ЧЕМПИОНАТА ФНТР ПО НАСТОЛЬНОМУ ТЕННИСУ СЕЗОНА 2016/2017 гг.")
        championship_label.setAlignment(Qt.AlignCenter)
        championship_label.setStyleSheet("font-weight: bold; font-size: 10pt;")
        layout.addWidget(championship_label)
        
        # Среди команд лиги
        league_label = QLabel("СРЕДИ КОМАНД ЛИГИ")
        league_label.setAlignment(Qt.AlignCenter)
        league_label.setStyleSheet("font-weight: bold; font-size: 10pt;")
        layout.addWidget(league_label)
        
        # Поля для ввода информации о матче
        info_widget = QWidget()
        info_layout = QHBoxLayout(info_widget)
        
        # Группа тур
        group_layout = QVBoxLayout()
        group_label = QLabel("группа тур")
        group_label.setAlignment(Qt.AlignCenter)
        group_layout.addWidget(group_label)
        self.group_input = QLineEdit()
        self.group_input.setPlaceholderText("Введите группу и тур")
        group_layout.addWidget(self.group_input)
        info_layout.addLayout(group_layout)
        
        # И
        and_label = QLabel("и")
        and_label.setAlignment(Qt.AlignCenter)
        and_label.setStyleSheet("font-weight: bold; font-size: 12pt;")
        info_layout.addWidget(and_label)
        
        # Команда А
        team_a_layout = QVBoxLayout()
        team_a_label = QLabel("команда")
        team_a_label.setAlignment(Qt.AlignCenter)
        team_a_layout.addWidget(team_a_label)
        self.team_a_input = QLineEdit()
        self.team_a_input.setPlaceholderText("Название команды А")
        team_a_layout.addWidget(self.team_a_input)
        info_layout.addLayout(team_a_layout)
        
        # Команда Б
        team_b_layout = QVBoxLayout()
        team_b_label = QLabel("команда")
        team_b_label.setAlignment(Qt.AlignCenter)
        team_b_layout.addWidget(team_b_label)
        self.team_b_input = QLineEdit()
        self.team_b_input.setPlaceholderText("Название команды Б")
        team_b_layout.addWidget(self.team_b_input)
        info_layout.addLayout(team_b_layout)
        
        layout.addWidget(info_widget)
        
        return header_widget
    
    def create_main_table(self):
        """Создание основной таблицы протокола"""
        table_widget = QTableWidget(5, 17)  # 5 строк, 17 колонок
        
        # Устанавливаем заголовки колонок
        headers = [
            "Фамилия Имя игроков", "П", "П1", "П2",
            "Фамилия Имя игроков", "П", "П1", "П2",
            "1-я партия", "2-я партия", "3-я партия", 
            "4-я партия", "5-я партия", "ОБЩИЙ СЧЁТ",
            "ОБЩИЙ СЧЁТ ПАРТИЙ", "КОМАНДНЫЕ", "ОЧКИ"
        ]
        table_widget.setHorizontalHeaderLabels(headers)
        
        # Устанавливаем заголовки строк
        row_labels = ["A", "B", "C", "A/D", "B/D"]
        table_widget.setVerticalHeaderLabels(row_labels)
        
        # Настройка таблицы
        table_widget.setShowGrid(True)
        table_widget.setGridStyle(Qt.SolidLine)
        table_widget.verticalHeader().setDefaultSectionSize(30)
        
        # Устанавливаем ширину колонок
        column_widths = [100, 30, 30, 30, 100, 30, 30, 30, 
                        50, 50, 50, 50, 50, 70, 70, 70, 50]
        for i, width in enumerate(column_widths):
            table_widget.setColumnWidth(i, width)
        
        # Делаем заголовки жирными
        font = table_widget.horizontalHeader().font()
        font.setBold(True)
        table_widget.horizontalHeader().setFont(font)
        table_widget.verticalHeader().setFont(font)
        
        # Добавляем объединенные ячейки для заголовков
        table_widget.setSpan(0, 13, 1, 2)  # Объединяем "ОБЩИЙ СЧЁТ" и "ОБЩИЙ СЧЁТ ПАРТИЙ"
        
        # Устанавливаем стиль
        table_widget.setStyleSheet("""
            QTableWidget {
                border: 2px solid #000000;
                gridline-color: #000000;
            }
            QHeaderView::section {
                background-color: #f0f0f0;
                border: 1px solid #000000;
                padding: 2px;
                font-weight: bold;
            }
        """)
        
        return table_widget
    
    def create_signatures_widget(self):
        """Создание виджета для подписей"""
        signatures_widget = QWidget()
        layout = QVBoxLayout(signatures_widget)
        layout.setSpacing(10)
        
        # Победитель
        winner_widget = QWidget()
        winner_layout = QHBoxLayout(winner_widget)
        winner_label = QLabel("ПОБЕДИТЕЛЬ:")
        winner_label.setStyleSheet("font-weight: bold; font-size: 11pt;")
        winner_layout.addWidget(winner_label)
        self.winner_input = QLineEdit()
        self.winner_input.setPlaceholderText("Укажите команду-победителя")
        winner_layout.addWidget(self.winner_input)
        layout.addWidget(winner_widget)
        
        # Подписи тренеров
        coaches_widget = QWidget()
        coaches_layout = QHBoxLayout(coaches_widget)
        
        # Тренер команды А
        coach_a_layout = QVBoxLayout()
        coach_a_label = QLabel("Подпись тренера-представителя команды А")
        coach_a_label.setStyleSheet("text-decoration: underline;")
        coach_a_layout.addWidget(coach_a_label)
        self.coach_a_input = QLineEdit()
        self.coach_a_input.setPlaceholderText("Фамилия И.О.")
        coach_a_layout.addWidget(self.coach_a_input)
        coaches_layout.addLayout(coach_a_layout)
        
        # Тренер команды Б
        coach_b_layout = QVBoxLayout()
        coach_b_label = QLabel("Подпись тренера-представителя команды Б")
        coach_b_label.setStyleSheet("text-decoration: underline;")
        coach_b_layout.addWidget(coach_b_label)
        self.coach_b_input = QLineEdit()
        self.coach_b_input.setPlaceholderText("Фамилия И.О.")
        coach_b_layout.addWidget(self.coach_b_input)
        coaches_layout.addLayout(coach_b_layout)
        
        layout.addWidget(coaches_widget)
        
        # Подпись судьи
        judge_widget = QWidget()
        judge_layout = QHBoxLayout(judge_widget)
        
        judge_label = QLabel("Фамилия Имя и подпись ведущего судьи")
        judge_label.setStyleSheet("text-decoration: underline;")
        judge_layout.addWidget(judge_label)
        
        self.judge_input = QLineEdit()
        self.judge_input.setPlaceholderText("Фамилия И.О. судьи")
        judge_layout.addWidget(self.judge_input)
        
        layout.addWidget(judge_widget)
        
        return signatures_widget
    
    def create_legend_widget(self):
        """Создание легенды сокращений"""
        legend_widget = QWidget()
        layout = QVBoxLayout(legend_widget)
        
        legend_label = QLabel("Сокращение в таблицах:")
        legend_label.setStyleSheet("font-weight: bold;")
        layout.addWidget(legend_label)
        
        # Таблица с сокращениями
        legend_table = QTableWidget(4, 2)
        legend_table.setHorizontalHeaderLabels(["Сокращение", "Расшифровка"])
        legend_table.verticalHeader().setVisible(False)
        
        legend_data = [
            ("Ж", "желтая карточка тренеру-представителю"),
            ("К", "красная карточка тренеру-представителю"),
            ("П", "предупреждение"),
            ("П1", "одно штрафное очко"),
            ("П2", "два штрафных очка")
        ]
        
        legend_table.setRowCount(len(legend_data))
        for i, (abbr, desc) in enumerate(legend_data):
            legend_table.setItem(i, 0, QTableWidgetItem(abbr))
            legend_table.setItem(i, 1, QTableWidgetItem(desc))
        
        legend_table.resizeColumnsToContents()
        legend_table.setMaximumHeight(150)
        
        layout.addWidget(legend_table)
        
        return legend_widget
    
    def fill_test_data(self):
        """Заполнение тестовыми данными"""
        for i, protocol in enumerate(self.protocol_widgets):
            # Заполняем информацию о матче
            if hasattr(self, 'group_input'):
                self.group_input.setText(f"Группа {i+1}, Тур {i+1}")
            
            if hasattr(self, 'team_a_input'):
                self.team_a_input.setText(f"Команда А-{i+1}")
            
            if hasattr(self, 'team_b_input'):
                self.team_b_input.setText(f"Команда Б-{i+1}")
            
            # Заполняем таблицу
            table = protocol.table_widget
            
            # Данные для заполнения
            players_a = ["Иванов И.", "Петров П.", "Сидоров С.", "Иванов/Смирнов", "Петров/Смирнов"]
            players_b = ["Кузнецов К.", "Михайлов М.", "Алексеев А.", "Кузнецов/Волков", "Михайлов/Волков"]
            
            for row in range(5):
                # Игроки команды А
                table.setItem(row, 0, QTableWidgetItem(players_a[row]))
                
                # Штрафы команды А
                if row == 1:
                    table.setItem(row, 1, QTableWidgetItem("П"))
                    table.setItem(row, 2, QTableWidgetItem("1"))
                
                # Игроки команды Б
                table.setItem(row, 4, QTableWidgetItem(players_b[row]))
                
                # Штрафы команды Б
                if row == 3:
                    table.setItem(row, 5, QTableWidgetItem("П"))
                    table.setItem(row, 6, QTableWidgetItem("1"))
                
                # Результаты партий
                scores = [
                    ["11:8", "11:9", "11:7", "", "", "3:0", "33:24", "1"],
                    ["8:11", "11:9", "9:11", "11:8", "8:11", "2:3", "47:50", "0"],
                    ["11:5", "11:6", "11:4", "", "", "3:0", "33:15", "1"],
                    ["11:9", "9:11", "11:6", "11:8", "", "3:1", "42:34", "1"],
                    ["7:11", "11:7", "8:11", "11:9", "9:11", "2:3", "46:49", "0"]
                ]
                
                for col in range(8, 17):
                    if col - 8 < len(scores[row]):
                        table.setItem(row, col, QTableWidgetItem(scores[row][col-8]))
            
            # Заполняем подписи
            if hasattr(self, 'winner_input'):
                self.winner_input.setText(f"Команда А-{i+1}")
            
            if hasattr(self, 'coach_a_input'):
                self.coach_a_input.setText(f"Тренеров А.И.")
            
            if hasattr(self, 'coach_b_input'):
                self.coach_b_input.setText(f"Тренеров Б.И.")
            
            if hasattr(self, 'judge_input'):
                self.judge_input.setText(f"Судьин С.П.")
    
    def clear_all(self):
        """Очистка всех полей"""
        for protocol in self.protocol_widgets:
            table = protocol.table_widget
            for row in range(table.rowCount()):
                for col in range(table.columnCount()):
                    table.setItem(row, col, None)
            
            # Очищаем поля ввода
            widgets = [
                self.group_input, self.team_a_input, self.team_b_input,
                self.winner_input, self.coach_a_input, self.coach_b_input,
                self.judge_input
            ]
            
            for widget in widgets:
                if hasattr(self, widget.__class__.__name__):
                    widget.clear()
    
    def generate_pdf(self):
        """Генерация PDF файла с 4 протоколами на листе А4"""
        filename, _ = QFileDialog.getSaveFileName(
            self, "Сохранить PDF", "", "PDF Files (*.pdf)"
        )
        
        if not filename:
            return
        
        # Создаем PDF документ
        doc = SimpleDocTemplate(
            filename,
            pagesize=A4,
            leftMargin=15*mm,
            rightMargin=15*mm,
            topMargin=10*mm,
            bottomMargin=10*mm
        )
        
        elements = []
        styles = getSampleStyleSheet()
        
        # Создаем стиль для заголовков
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Normal'],
            fontSize=10,
            alignment=1,  # Центр
            spaceAfter=6,
            spaceBefore=6
        )
        
        # Добавляем 4 протокола
        for i, protocol in enumerate(self.protocol_widgets):
            # Заголовок протокола
            elements.append(Paragraph(f"ПРОТОКОЛ КОМАНДНОГО МАТЧА #{i+1}", header_style))
            elements.append(Spacer(1, 5*mm))
            
            # Шапка документа
            elements.append(Paragraph(
                "Министерство спорта Российской Федерации / Федерация настольного тенниса России",
                header_style
            ))
            elements.append(Paragraph(
                "Директорат Командного Чемпионата ФНТР",
                header_style
            ))
            
            elements.append(Spacer(1, 2*mm))
            
            # Информация о матче
            info_text = f"""
            Группа, тур: {self.group_input.text() if hasattr(self, 'group_input') else ''}<br/>
            Команда А: {self.team_a_input.text() if hasattr(self, 'team_a_input') else ''}<br/>
            Команда Б: {self.team_b_input.text() if hasattr(self, 'team_b_input') else ''}
            """
            elements.append(Paragraph(info_text, styles['Normal']))
            
            elements.append(Spacer(1, 5*mm))
            
            # Создаем таблицу данных
            table = protocol.table_widget
            data = []
            
            # Заголовки таблицы
            headers = [
                "Фамилия Имя игроков", "П", "П1", "П2",
                "Фамилия Имя игроков", "П", "П1", "П2",
                "1-я партия", "2-я партия", "3-я партия", 
                "4-я партия", "5-я партия", "ОБЩИЙ СЧЁТ",
                "ОБЩИЙ СЧЁТ ПАРТИЙ", "КОМАНДНЫЕ", "ОЧКИ"
            ]
            data.append(headers)
            
            # Данные таблицы
            for row in range(table.rowCount()):
                row_data = []
                for col in range(table.columnCount()):
                    item = table.item(row, col)
                    row_data.append(item.text() if item else "")
                data.append(row_data)
            
            # Создаем таблицу для PDF
            pdf_table = Table(data, colWidths=[60, 15, 15, 15, 60, 15, 15, 15, 
                                              25, 25, 25, 25, 25, 35, 35, 35, 25])
            
            # Стиль таблицы
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
            ])
            
            pdf_table.setStyle(table_style)
            elements.append(pdf_table)
            
            elements.append(Spacer(1, 5*mm))
            
            # Подписи
            signatures = [
                f"ПОБЕДИТЕЛЬ: {self.winner_input.text() if hasattr(self, 'winner_input') else ''}",
                f"Тренер команды А: {self.coach_a_input.text() if hasattr(self, 'coach_a_input') else ''}",
                f"Тренер команды Б: {self.coach_b_input.text() if hasattr(self, 'coach_b_input') else ''}",
                f"Судья: {self.judge_input.text() if hasattr(self, 'judge_input') else ''}"
            ]
            
            for signature in signatures:
                elements.append(Paragraph(signature, styles['Normal']))
            
            elements.append(Spacer(1, 5*mm))
            
            # Легенда
            elements.append(Paragraph("Сокращение в таблицах:", styles['Normal']))
            legend = [
                "Ж - желтая карточка тренеру-представителю",
                "К - красная карточка тренеру-представителю", 
                "П - предупреждение",
                "П1 - одно штрафное очко",
                "П2 - два штрафных очка"
            ]
            
            for item in legend:
                elements.append(Paragraph(item, styles['Normal']))
            
            # Добавляем разрыв страницы между протоколами
            if i < 3:
                elements.append(PageBreak())
        
        # Генерируем PDF
        doc.build(elements)
        
        QMessageBox.information(self, "Успех", f"PDF файл сохранен:\n{filename}")
    
    def print_protocols(self):
        """Печать протоколов"""
        printer = QPrinter(QPrinter.HighResolution)
        printer.setPageSize(QPrinter.A4)
        printer.setPageOrientation(QPrinter.Portrait)
        
        dialog = QPrintDialog(printer, self)
        if dialog.exec_() == QPrintDialog.Accepted:
            # Создаем документ для печати
            painter = QPainter(printer)
            
            # Рассчитываем размеры для 4 протоколов на листе
            page_rect = printer.pageRect()
            width = page_rect.width()
            height = page_rect.height()
            
            # Размеры для одного протокола
            proto_width = width
            proto_height = height / 4
            
            # Печатаем каждый протокол
            for i in range(4):
                if i > 0:
                    printer.newPage()
                
                # Устанавливаем область для текущего протокола
                painter.setViewport(0, int(i * proto_height), 
                                   int(proto_width), int(proto_height))
                painter.setWindow(0, 0, 800, 600)
                
                # Рисуем протокол
                self.draw_protocol(painter, i)
            
            painter.end()
    
    def draw_protocol(self, painter, protocol_index):
        """Рисует один протокол на QPainter"""
        painter.save()
        
        # Задаем фон
        painter.fillRect(0, 0, 800, 600, QColor(255, 255, 255))
        
        # Заголовок
        painter.setFont(QFont('DejaVuSerif-Italic', 14, QFont.Bold))
        painter.drawText(QRect(0, 10, 800, 30), 
                        Qt.AlignCenter, 
                        f"ПРОТОКОЛ КОМАНДНОГО МАТЧА #{protocol_index + 1}")
        
        # Министерство и Федерация
        painter.setFont(QFont('DejaVuSerif-Italic', 10, QFont.Bold))
        painter.drawText(QRect(0, 50, 800, 20), 
                        Qt.AlignCenter, 
                        "Министерство спорта Российской Федерации / Федерация настольного тенниса России")
        
        # Директорат
        painter.drawText(QRect(0, 70, 800, 20), 
                        Qt.AlignCenter, 
                        "Директорат Командного Чемпионата ФНТР")
        
        # Основная таблица
        painter.setPen(QPen(QColor(0, 0, 0), 1))
        
        # Рисуем таблицу
        start_x = 20
        start_y = 120
        cell_width = 40
        cell_height = 25
        
        # Рисуем заголовки колонок
        headers = ["Игрок А", "П", "П1", "П2", "Игрок Б", "П", "П1", "П2",
                  "1-я", "2-я", "3-я", "4-я", "5-я", "Общий", "Партии", "Ком.", "Очки"]
        
        for i, header in enumerate(headers):
            rect = QRect(start_x + i * cell_width, start_y, cell_width, cell_height)
            painter.drawRect(rect)
            painter.drawText(rect, Qt.AlignCenter, header)
        
        # Рисуем строки
        for row in range(5):
            row_y = start_y + (row + 1) * cell_height
            for col in range(17):
                rect = QRect(start_x + col * cell_width, row_y, cell_width, cell_height)
                painter.drawRect(rect)
        
        # Подписи
        painter.setFont(QFont('DejaVuSerif-Italic', 10))
        painter.drawText(QRect(20, 300, 400, 30), 
                        "Тренер команды А: __________________________")
        painter.drawText(QRect(20, 340, 400, 30), 
                        "Тренер команды Б: __________________________")
        painter.drawText(QRect(20, 380, 400, 30), 
                        "Судья: ____________________________________")
        
        painter.restore()
    
    def save_to_word(self):
        """Сохранение в Word (альтернатива PDF)"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            
            filename, _ = QFileDialog.getSaveFileName(
                self, "Сохранить Word документ", "", "Word Files (*.docx)"
            )
            
            if not filename:
                return
            
            doc = Document()
            
            # Настройка страницы
            section = doc.sections[0]
            section.page_height = Inches(11.69)
            section.page_width = Inches(8.27)
            
            for i, protocol in enumerate(self.protocol_widgets):
                # Добавляем заголовок
                doc.add_heading(f'ПРОТОКОЛ КОМАНДНОГО МАТЧА #{i+1}', 1)
                
                # Добавляем информацию
                doc.add_paragraph(
                    f"Группа, тур: {self.group_input.text() if hasattr(self, 'group_input') else ''}"
                )
                doc.add_paragraph(
                    f"Команда А: {self.team_a_input.text() if hasattr(self, 'team_a_input') else ''}"
                )
                doc.add_paragraph(
                    f"Команда Б: {self.team_b_input.text() if hasattr(self, 'team_b_input') else ''}"
                )
                
                # Добавляем разрыв страницы между протоколами
                if i < 3:
                    doc.add_page_break()
            
            doc.save(filename)
            QMessageBox.information(self, "Успех", f"Word документ сохранен:\n{filename}")
            
        except ImportError:
            QMessageBox.warning(self, "Ошибка", 
                              "Для сохранения в Word установите библиотеку: pip install python-docx")
            
    def create_exact_table_widget(self):
        """Создание точной копии таблицы из Word документа"""
        table_widget = QTableWidget(6, 17)  # 6 строк (1 заголовок + 5 данных)
        
        # Настраиваем объединенные ячейки как в Word
        table_widget.setSpan(0, 13, 1, 2)  # Объединяем "ОБЩИЙ СЧЁТ ПАРТИЙ"
        
        # Заголовки как в Word файле
        headers = [
            "Фамилия Имя игроков", "П", "П1", "П2",
            "Фамилия Имя игроков", "П", "П1", "P2",
            "1-я партия", "2-я партия", "3-я партия", 
            "4-я партия", "5-я партия", "ОБЩИЙ СЧЁТ",
            "ОБЩИЙ СЧЁТ ПАРТИЙ", "КОМАНДНЫЕ", "ОЧКИ"
        ]
        
        # Устанавливаем двойные линии для границ как в Word
        table_widget.setStyleSheet("""
            QTableWidget {
                border: 2px solid black;
            }
            QTableWidget::item {
                border: 1px solid black;
                padding: 2px;
            }
            QHeaderView::section {
                background-color: #f0f0f0;
                border: 1px solid black;
                border-left: 2px solid black;
                border-right: 2px solid black;
                font-weight: bold;
                padding: 2px;
            }
        """)
        
        # Устанавливаем фиксированные размеры как в Word
        table_widget.verticalHeader().setVisible(False)
        
        # Заполняем данные строк
        row_labels = ["A", "B", "C", "A/D", "B/D"]
        
        for row in range(1, 6):
            # Добавляем маркеры строк слева
            row_header = QTableWidgetItem(row_labels[row-1])
            row_header.setTextAlignment(Qt.AlignCenter)
            row_header.setFont(QFont('DejaVuSerif-Italic', 9, QFont.Bold))
            table_widget.setVerticalHeaderItem(row-1, row_header)
        
        return table_widget

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    
    # Устанавливаем русскую локаль
    translator = QTranslator()
    translator.load("qt_ru")
    app.installTranslator(translator)
    
    window = ProtocolGenerator()
    window.show()
    
    sys.exit(app.exec_())